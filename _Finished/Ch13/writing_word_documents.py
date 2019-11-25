import docx
import os
os.chdir(os.path.dirname(os.path.realpath(__file__)))

def create_simple():
    doc = docx.Document()
    doc.add_paragraph("Howdy world!")
    doc.save("Workfiles/helloworld.docx")

def create_simple_multiple_paras():
    doc = docx.Document()
    para_obj1 = doc.add_paragraph("1st paragraph!")
    para_obj2 = doc.add_paragraph("3rd paragraph!")
    para_obj3 = doc.add_paragraph("1st paragraph again!")
    para_obj4 = doc.add_paragraph("2nd paragraph!")
    para_obj4.add_run("\nMan... The ordering is all fucked.")
    doc.save("Workfiles/multipleParagraphs.docx")

def adding_headings():
    doc = docx.Document()
    doc.add_heading('Header 0', 0)
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)
    doc.save("Workfiles/addingHeadings.docx")

def adding_line_and_page_breaks():
    doc = docx.Document()
    doc.add_paragraph("This is the first page!")
    doc.add_page_break() # NEW METHOD 
    doc.add_paragraph("This is the second page!")
    doc.save("Workfiles/AddingLinesAndPageBreaks.docx")

def adding_pictures():
    doc = docx.Document()
    doc.add_picture('Workfiles/zophie.png', width=docx.shared.Cm(2.52), height=docx.shared.Cm(4))
    # doc.add_picture('Workfiles/zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
    doc.save("Workfiles/Picture.docx")

adding_pictures()