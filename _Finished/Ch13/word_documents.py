import os
import docx

### Changing CWD to this folder.
os.chdir(os.path.dirname(os.path.realpath(__file__)))

doc = docx.Document('Workfiles/demo.docx')
print(len(doc.paragraphs))  # Output: 7
print(doc.paragraphs[0].text) # Document Title
print(doc.paragraphs[1].text) # A plain paragraph with some bold and some italic
print(len(doc.paragraphs[1].runs)) # 5
print(doc.paragraphs[1].runs[0].text) # A plain paragraph with
print(doc.paragraphs[1].runs[1].text) #  some 
print(doc.paragraphs[1].runs[2].text) # 'bold'
print(doc.paragraphs[1].runs[3].text) # ' and some '
print(doc.paragraphs[1].runs[4].text) # 'italic'