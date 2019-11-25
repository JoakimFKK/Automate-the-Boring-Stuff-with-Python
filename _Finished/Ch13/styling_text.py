import docx
import os
os.chdir(os.path.dirname(os.path.realpath(__file__)))

doc = docx.Document("Workfiles/demo.docx")
print(doc.paragraphs[0].text) # Document Title
print(doc.paragraphs[0].style) # _ParagraphStyle('Title') id: 2513206750704
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text) # A plain paragraph with some bold and some italic
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text,
       doc.paragraphs[1].runs[3].text, doc.paragraphs[1].runs[4].text)) # ('A plain paragraph with', ' some ', 'bold', ' and some ', 'italic')
doc.paragraphs[1].runs[0].style = "Quote Char" # Spaces are allowed in this version.
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('Workfiles/restyled.docx')